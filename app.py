from flask import Flask, render_template, request, send_from_directory
from docx import Document
from flask import send_file
from docx2pdf import convert
import time
import openai
import re
import os

app = Flask(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(BASE_DIR, 'static')


@app.route('/')
def index():
    try:
        if os.path.exists('static/output.docx'):
            os.remove('static/output.docx')
            os.remove('static/output.pdf')
            return render_template('index.html')

        else:
            return render_template('index.html')

    except Exception as e:
        return str(e)


@app.route('/form')
def form():
    try:
        if os.path.exists('static/output.docx'):
            os.remove('static/output.docx')
            os.remove('static/output.pdf')
            return render_template('form.html')

        else:
            return render_template('form.html')
    except Exception as e:
        return str(e)


@app.route('/generate', methods=['POST'])
def generate():
    try:
        title = request.form['title']
        author = request.form['author']
        problemStatement = request.form['problem']
        proposedSl = request.form['solution']
        result = request.form['result']

        # Generate the technical review paper using OpenAI API
        openai.api_key = "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"

        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo", messages=[{"role": "user", "content": "Write Abstract, only Introduction, Keywords, Conclusion about "+title+" in detail. Give me in this format Abstract:, Introduction:, Keywords: & Conclusion:."}])

        part1 = completion.choices[0].message.content

        # Define regular expressions for the abstract, introduction, keywords, and conclusion
        abstract_re = 'Abstract:(.*?)Introduction:'
        introduction_re = 'Introduction:(.*?)Keywords:'
        keywords_re = 'Keywords:(.*?)Conclusion:'
        conclusion_re = 'Conclusion:(.*?)$'

        # Extract the abstract, introduction, keywords, and conclusion using the regular expressions
        abstract_match = re.search(abstract_re, part1, re.DOTALL)
        introduction_match = re.search(introduction_re, part1, re.DOTALL)
        keywords_match = re.search(keywords_re, part1, re.DOTALL)
        conclusion_match = re.search(conclusion_re, part1, re.DOTALL)

        # Check if a match was found for each section
        if abstract_match:
            abstract = abstract_match.group(1).strip()
        else:
            abstract = "No Abstarct found in the response."

        if introduction_match:
            introduction = introduction_match.group(1).strip()
        else:
            introduction = "No Introduction found in the response."

        if keywords_match:
            keywords = keywords_match.group(1).strip()
        else:
            keywords = "No Keywords found in the response."

        if conclusion_match:
            conclusion = conclusion_match.group(1).strip()
        else:
            conclusion = "No Conclusion found in the response."

        time.sleep(10)

        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo", messages=[{"role": "user", "content": "write 5 references(List references numbering in square parenthesis format) about "+title+". write detailed Literature survey according to generated references.. give me in this format References: & Literature Survey: "}])
        part2 = completion.choices[0].message.content

        # Define regular expressions for the references and literature survey
        reference_re = 'References:(.*?)Literature Survey:'
        literature_survey_re = 'Literature Survey:(.*?)$'

        # Extract the references and literature survey using the regular expressions
        reference_match = re.search(reference_re, part2, re.DOTALL)
        literature_survey_match = re.search(
            literature_survey_re, part2, re.DOTALL)

        if literature_survey_match:
            literature_survey = literature_survey_match.group(1).strip()
        else:
            literature_survey = "No literature survey found in the response."

        if reference_match:
            references = reference_match.group(1).strip()
        else:
            references = "No References found in the response."

        # Print the extracted sections
        document = Document('IEEE.docx')  # Open the document

        # Access the title paragraph and change its text
        title_paragraph = document.paragraphs[0]
        title_paragraph.text = title.upper()

        title_paragraph = document.paragraphs[1]
        title_paragraph.text = author.title()

        title_paragraph = document.paragraphs[5]
        title_paragraph.text = "Abstract- "+abstract

        title_paragraph = document.paragraphs[6]
        title_paragraph.text = 'Keywords- '+keywords.title()

        title_paragraph = document.paragraphs[8]
        title_paragraph.text = introduction

        title_paragraph = document.paragraphs[10]
        title_paragraph.text = literature_survey

        title_paragraph = document.paragraphs[12]
        title_paragraph.text = problemStatement

        title_paragraph = document.paragraphs[14]
        title_paragraph.text = proposedSl

        title_paragraph = document.paragraphs[16]
        title_paragraph.text = result

        title_paragraph = document.paragraphs[18]
        title_paragraph.text = conclusion

        title_paragraph = document.paragraphs[20]
        title_paragraph.text = "I extend my sincere gratitude to all the individuals who have played a pivotal role in making this mini project on "+title.title()+" a success. I am immensely grateful to my mentor, (mentor's name), whose guidance, encouragement, and support have been invaluable in every step of this project. Without their expertise and timely assistance, the completion of this project would not have been possible.I also thank my team members, (names of team members), for their relentless hard work, dedication, and cooperation throughout the project. Their inputs and insights have significantly contributed to the quality of the project.I am thankful to the management of(name of the institution) for providing us with all the necessary resources such as space, equipment, software, and infrastructure, which were essential in accomplishing this project.In conclusion, I express my sincere appreciation to all who have contributed to the successful completion of this project. I hope this project on "+title.title()+" will be useful for the advancement of parking management and help reduce traffic congestion in urban areas."

        title_paragraph = document.paragraphs[22]
        title_paragraph.text = references

        document.save('static/output.docx')  # Save the modified document
        convert('static/output.docx', 'static/output.pdf')

        # Return the generated paper to the user
        return render_template('result.html')

    except Exception as e:
        return str(e)


@app.route('/download')
def download():
    try:
        document = Document()
        # Add content to the document here
        return send_from_directory(STATIC_DIR, "output.docx", as_attachment=True)
    except Exception as e:
        return str(e)


if __name__ == '__main__':
    app.run(port=3999,debug=True)
